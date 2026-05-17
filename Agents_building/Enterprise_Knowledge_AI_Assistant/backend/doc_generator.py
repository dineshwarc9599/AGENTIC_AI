from docx import Document
import uuid
import os

OUTPUT_DIR = "output_docs"

os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_docx(content):
    doc = Document()

    lines = content.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line.endswith(":"):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    filename = f"generated_{uuid.uuid4().hex}.docx"

    filepath = os.path.join(OUTPUT_DIR, filename)

    doc.save(filepath)

    return filepath, filename